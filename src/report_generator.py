#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
报告生成模块，用于生成Word文档格式的报告
"""

import os
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config


def create_report():
    """
    创建一个新的报告文档

    Returns:
        docx.Document: 创建的文档对象
    """
    # 创建文档对象
    doc = Document()

    # 设置文档标题
    title = doc.add_heading(config.REPORT_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加文档属性
    doc.core_properties.author = config.REPORT_AUTHOR
    doc.core_properties.title = config.REPORT_TITLE

    return doc


def set_cell_background(cell, color):
    """
    设置单元格背景颜色

    Args:
        cell: 单元格对象
        color: RGB颜色元组，例如(255, 255, 0)表示黄色
    """
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), f"{color[0]:02x}{color[1]:02x}{color[2]:02x}")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_image_pair_to_report(
    doc, original_image_path, corrected_image_path, description, action, location=""
):
    """
    将一对图片及其描述和纠正措施添加到报告中，格式与示例一致

    Args:
        doc (docx.Document): 文档对象
        original_image_path (str): 原始图片路径
        corrected_image_path (str): 纠正后的图片路径
        description (str): 描述
        action (str): 纠正措施
        location (str, optional): 位置信息，默认为空

    Returns:
        docx.Document: 更新后的文档对象
    """
    # 创建一个1行1列的表格作为主表格
    main_table = doc.add_table(rows=5, cols=1)
    main_table.style = "Table Grid"

    # 设置第一行为"Observation"标题行，并设置黄色背景
    header_cell = main_table.cell(0, 0)
    header_cell.text = "Observation"
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.bold = True
    header_run.font.size = Pt(14)
    # 设置黄色背景
    set_cell_background(header_cell, (255, 255, 0))  # 黄色

    # 设置第二行为位置信息
    location_cell = main_table.cell(1, 0)
    location_cell.text = f"Location: {location}"
    location_para = location_cell.paragraphs[0]
    location_run = location_para.runs[0]
    location_run.bold = True

    # 第三行为图片表格
    images_cell = main_table.cell(2, 0)
    # 在单元格中创建一个嵌套表格用于放置图片
    images_table = images_cell.add_table(rows=2, cols=2)

    # 设置"Before"和"After"标题
    before_cell = images_table.cell(0, 0)
    before_cell.text = "Before"
    before_para = before_cell.paragraphs[0]
    before_run = before_para.runs[0]
    before_run.bold = True
    before_run.font.size = Pt(12)

    after_cell = images_table.cell(0, 1)
    after_cell.text = "After"
    after_para = after_cell.paragraphs[0]
    after_run = after_para.runs[0]
    after_run.bold = True
    after_run.font.size = Pt(12)

    # 添加图片
    before_img_cell = images_table.cell(1, 0)
    before_img_para = before_img_cell.paragraphs[0]
    before_img_run = before_img_para.add_run()
    before_img_run.add_picture(original_image_path, width=Inches(3.0))

    after_img_cell = images_table.cell(1, 1)
    after_img_para = after_img_cell.paragraphs[0]
    after_img_run = after_img_para.add_run()
    after_img_run.add_picture(corrected_image_path, width=Inches(3.0))

    # 第四行为描述和纠正措施表格
    desc_cell = main_table.cell(3, 0)
    # 在单元格中创建一个嵌套表格用于放置描述和纠正措施
    desc_table = desc_cell.add_table(rows=2, cols=2)

    # 设置"Description"和"Corrective Action"标题
    desc_header_cell = desc_table.cell(0, 0)
    desc_header_cell.text = "Description"
    desc_header_para = desc_header_cell.paragraphs[0]
    desc_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_header_run = desc_header_para.runs[0]
    desc_header_run.bold = True

    action_header_cell = desc_table.cell(0, 1)
    action_header_cell.text = "Corrective Action"
    action_header_para = action_header_cell.paragraphs[0]
    action_header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    action_header_run = action_header_para.runs[0]
    action_header_run.bold = True

    # 添加描述和纠正措施内容
    desc_content_cell = desc_table.cell(1, 0)
    desc_content_cell.text = description
    desc_content_para = desc_content_cell.paragraphs[0]
    desc_content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    action_content_cell = desc_table.cell(1, 1)
    action_content_cell.text = action
    action_content_para = action_content_cell.paragraphs[0]
    action_content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第五行为空行，用于分隔不同的观察项
    main_table.cell(4, 0).text = ""

    return doc


def add_image_pair_to_template_page(
    doc,
    page_index,
    original_image_path,
    corrected_image_path,
    description,
    action,
    location="",
):
    """
    将一对图片及其描述和纠正措施添加到模板页面中

    Args:
        doc (docx.Document): 文档对象
        page_index (int): 页面索引
        original_image_path (str): 原始图片路径
        corrected_image_path (str): 纠正后的图片路径
        description (str): 描述
        action (str): 纠正措施
        location (str, optional): 位置信息，默认为空

    Returns:
        docx.Document: 更新后的文档对象
    """
    # 查找表格
    tables = doc.tables
    if page_index >= len(tables):
        print(
            f"警告：页面索引 {page_index} 超出范围，当前文档只有 {len(tables)} 个表格"
        )
        return doc

    # 获取当前页面的表格
    table = tables[page_index]

    # 设置位置信息
    # 假设位置信息在第二行
    location_cell = None
    for row in table.rows:
        for cell in row.cells:
            if cell.text.startswith("Location:"):
                location_cell = cell
                break
        if location_cell:
            break

    if location_cell:
        location_cell.text = f"Location: {location}"
        for paragraph in location_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    else:
        print("警告：未找到位置信息单元格")

    # 查找Before和After图片单元格
    before_img_cell = None
    after_img_cell = None
    for row in table.rows:
        for cell in row.cells:
            if "Before" in cell.text:
                # 假设图片单元格在"Before"标题下方
                before_img_cell = table.cell(
                    row._tr.index(row._tr) + 1, cell._tc.index(cell._tc)
                )
            elif "After" in cell.text:
                # 假设图片单元格在"After"标题下方
                after_img_cell = table.cell(
                    row._tr.index(row._tr) + 1, cell._tc.index(cell._tc)
                )

    # 添加图片
    if before_img_cell:
        before_img_cell.text = ""  # 清除单元格内容
        before_img_para = before_img_cell.paragraphs[0]
        before_img_run = before_img_para.add_run()
        before_img_run.add_picture(original_image_path, width=Inches(3.0))
    else:
        print("警告：未找到Before图片单元格")

    if after_img_cell:
        after_img_cell.text = ""  # 清除单元格内容
        after_img_para = after_img_cell.paragraphs[0]
        after_img_run = after_img_para.add_run()
        after_img_run.add_picture(corrected_image_path, width=Inches(3.0))
    else:
        print("警告：未找到After图片单元格")

    # 查找Description和Corrective Action单元格
    desc_cell = None
    action_cell = None
    for row in table.rows:
        for cell in row.cells:
            if "Description" in cell.text:
                # 假设内容单元格在标题下方
                desc_cell = table.cell(
                    row._tr.index(row._tr) + 1, cell._tc.index(cell._tc)
                )
            elif "Corrective Action" in cell.text:
                # 假设内容单元格在标题下方
                action_cell = table.cell(
                    row._tr.index(row._tr) + 1, cell._tc.index(cell._tc)
                )

    # 添加描述和纠正措施内容
    if desc_cell:
        desc_cell.text = description
        for paragraph in desc_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("警告：未找到Description单元格")

    if action_cell:
        action_cell.text = action
        for paragraph in action_cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        print("警告：未找到Corrective Action单元格")

    return doc


def save_report(doc, output_path=None):
    """
    保存报告到指定路径

    Args:
        doc (docx.Document): 文档对象
        output_path (str, optional): 输出路径，如果为None则使用配置中的路径

    Returns:
        str: 保存的文件路径
    """
    if output_path is None:
        output_path = config.OUTPUT_REPORT

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        # 保存文档
        doc.save(output_path)
        print(f"报告已保存到: {output_path}")
        return output_path

    except Exception as e:
        print(f"保存报告时出错: {e}")
        return None


def generate_report(image_pairs_with_data, locations=None):
    """
    生成报告

    Args:
        image_pairs_with_data (list): 图片对及其数据的列表，每个元素是一个元组
                                     (原始图片路径, 纠正后的图片路径, 描述, 纠正措施)
        locations (list, optional): 位置信息列表，与image_pairs_with_data一一对应，默认为None

    Returns:
        str: 生成的报告路径
    """
    # 创建报告
    doc = create_report()

    # 添加图片对和描述
    for i, (original_image, corrected_image, description, action) in enumerate(
        image_pairs_with_data
    ):
        # 获取位置信息，如果没有提供则使用空字符串
        location = ""
        if locations and i < len(locations):
            location = locations[i]

        add_image_pair_to_report(
            doc, original_image, corrected_image, description, action, location
        )

    # 保存报告
    return save_report(doc)



def generate_report_from_template(
    image_pairs_with_data, locations=None, template_path=None
):
    """
    使用模板生成报告，为每个图像对复制模板页面

    Args:
        image_pairs_with_data (list): 图片对及其数据的列表，每个元素是一个元组
                                     (原始图片路径, 纠正后的图片路径, 描述, 纠正措施)
        locations (list, optional): 位置信息列表，与image_pairs_with_data一一对应，默认为None
        template_path (str, optional): 模板文件路径，如果为None则使用默认模板

    Returns:
        str: 生成的报告路径
    """
    # 如果没有指定模板路径，则使用默认模板
    if template_path is None:
        template_path = os.path.join(
            config.BASE_DIR, "template", "report-template.docx"
        )

    # 检查模板文件是否存在
    if not os.path.exists(template_path):
        print(f"错误：模板文件 {template_path} 不存在")
        return None

    try:
        # 使用更可靠的docxtpl库替代直接操作xml（需要安装）
        from docxtpl import DocxTemplate, InlineImage
        import jinja2
        from docx.shared import Inches

        # 创建输出文件路径
        output_path = config.OUTPUT_REPORT

        # 正确打开模板文件（移除with语句，直接使用DocxTemplate）
        doc = DocxTemplate(template_path)  # 直接传递文件路径
            
        # 准备上下文数据（添加图片路径处理）
        context = {
            'observations': [
                {
                    'location': locations[i] if locations and i < len(locations) else "",
                    # 使用docxtpl的Inches转换图片尺寸
                    'original_image': InlineImage(doc, original_image, width=Inches(3.0)),
                    'corrected_image': InlineImage(doc, corrected_image, width=Inches(3.0)),
                    'description': description,
                    'action': action
                } 
                for i, (original_image, corrected_image, description, action) 
                in enumerate(image_pairs_with_data)
            ]
        }

        # 渲染文档
        doc.render(context)

        # 保存文档
        doc.save(output_path)
        print(f"报告已保存到: {output_path}")
        return output_path

    except Exception as e:
        print(f"生成报告时出错: {e}")
        return None
